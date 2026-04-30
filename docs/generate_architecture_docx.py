"""Generate docs/TAL_Architecture.docx from the current architecture."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).parent / "TAOS_Architecture_v1.1.docx"

doc = Document()

# ---------- Styles ----------
styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

def h1(text):
    p = doc.add_heading(text, level=1)
    return p

def h2(text):
    return doc.add_heading(text, level=2)

def h3(text):
    return doc.add_heading(text, level=3)

def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    return p

def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v
    doc.add_paragraph()

# ==========================================================
# Title page
# ==========================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("Talent-Augmenting OS (TAOS)")
tr.bold = True
tr.font.size = Pt(28)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("System Architecture & Executive Overview")
sr.italic = True
sr.font.size = Pt(16)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Version 1.1  ·  30 April 2026  ·  Angelo Leone").font.size = Pt(11)

doc.add_page_break()

# ==========================================================
# SECTION 1: Executive Summary (non-technical)
# ==========================================================
h1("1. Executive Summary")

para(
    "The Talent-Augmenting OS (TAOS) is a personalised AI augmentation system that "
    "makes workers better at their jobs instead of dependent on AI. It sits between the "
    "worker and whichever AI assistant they use (ChatGPT, Claude, Gemini, Copilot, or any "
    "other) and changes how that AI behaves based on the worker's expertise, goals, and "
    "growth areas."
)

h2("The problem we solve")
para(
    "Today's AI tools have one setting: maximum helpfulness. That sounds good, but three "
    "predictable failures emerge in organisations that deploy them widely:"
)
bullet("De-skilling: workers lose the skills they stop practicing. In one 2024 clinical study, doctors performed worse after three months of AI use than before.")
bullet("Over-reliance: workers accept AI output without critical evaluation, so the human-plus-AI team performs worse than the AI alone (Buçinca, Harvard, 2021).")
bullet("Autopilot: junior employees who 'just hand in' AI output perform worse than those who engage critically (Mollick, Wharton, 2023).")
para(
    "TAOS is designed explicitly to prevent all three: not by limiting AI, but by "
    "making it smarter about when to help, when to coach, when to challenge, and when to "
    "step back."
)

h2("How it works in one paragraph")
para(
    "A 15-minute conversational assessment creates a personalised profile (a portable "
    "markdown file). That profile describes the worker's expertise in each of their key "
    "domains, what they are trying to grow, and which skills are at risk of atrophying. "
    "Every future AI interaction is then calibrated against that profile. In an expert "
    "domain the AI accelerates the worker; in a growth area it scaffolds and coaches; in "
    "an at-risk area it adds friction so the worker does the thinking themselves. The "
    "profile follows the worker across every AI tool they use."
)

h2("Why this matters for executives")
table(
    ["Outcome", "What it means for the business"],
    [
        ["Higher retention of skills", "Workforce capability compounds over time rather than eroding under AI use."],
        ["Lower over-reliance risk", "Fewer quality incidents caused by AI outputs being accepted unchallenged."],
        ["Measurable skill progression", "Organisation-level dashboard shows which skills are growing, stable, or at risk."],
        ["Platform-agnostic deployment", "No vendor lock-in: works with any LLM the company already pays for."],
        ["Research-backed", "Grounded in peer-reviewed work from Harvard, MIT, and Wharton: defensible to regulators, boards, and workers."],
    ],
)

h2("What we ship today")
bullet("A universal system prompt (2-minute install) that works with any LLM.")
bullet("Pre-configured instances for ChatGPT, Gemini, and Claude.")
bullet("A Model Context Protocol (MCP) server with 14 specialised tools for Claude Code, Cursor, Windsurf, and Claude Desktop.")
bullet("A 1-click Claude Desktop Extension and a Claude Cowork plugin.")
bullet("A hosted web application with Google sign-in, persistent profiles, and a remote MCP endpoint.")
bullet("An organisation-level analytics dashboard and an A/B testing framework.")

doc.add_page_break()

# ==========================================================
# SECTION 2: Architecture Overview
# ==========================================================
h1("2. Architecture Overview")

para(
    "TAOS is built as a layer, not a single product. The same underlying logic: "
    "assessment, scoring, task classification, and interaction tracking: is exposed "
    "through four tiers, each appropriate for a different level of integration and "
    "control."
)

h2("2.1 The four tiers")
table(
    ["Tier", "Surface", "Who it is for", "Setup"],
    [
        ["1", "Universal system prompt for any LLM", "Individuals, zero-dependency users", "2 min copy-paste"],
        ["2", "Custom GPT / Gemini Gem / Claude Project", "Small teams on a single LLM platform", "5 min"],
        ["3", "MCP server (stdio, Desktop Extension, Cowork plugin, Remote MCP)", "Power users on Claude Code, Cursor, Windsurf, Claude Desktop", "1-click to 10 min"],
        ["4", "Hosted web application", "Organisations wanting centralised profiles and dashboards", "Docker deploy"],
    ],
)

h2("2.2 The shared core")
para(
    "Every tier runs against the same portable profile format and the same behavioural "
    "rules (the TAOS system prompt in CLAUDE.md). This is the key architectural "
    "decision: by standardising the profile as markdown and the behaviour as a prompt, "
    "the worker's calibration travels with them across every platform change."
)

h2("2.3 System diagram (logical)")
para(
    "The diagram below shows the logical flow from user entry points through LLM "
    "orchestration, the TAOS core logic, transports, and storage. A rendered Mermaid "
    "version is maintained in docs/ARCHITECTURE.md."
)

para("User Entry Points", bold=True)
bullet("Any LLM (ChatGPT, Gemini, Claude web) via the universal system prompt.")
bullet("Platform-native: Custom GPT, Gemini Gem, Claude Project.")
bullet("Claude Code / Cursor / Windsurf via slash commands and the local MCP server.")
bullet("Claude Desktop via the 1-click Desktop Extension (.mcpb).")
bullet("Claude Code via the Claude Cowork plugin marketplace.")
bullet("Hosted web app at proworker-hosted.onrender.com (Google OAuth).")

para("LLM Orchestration", bold=True)
bullet("Claude (Code / Desktop): primary development surface.")
bullet("Gemini 2.5 Flash-Lite: powers the hosted conversational assessment.")
bullet("Any LLM: via the universal system prompt in custom instructions.")

para("TAOS Core Logic", bold=True)
bullet("MCP Server (mcp-server/src/server.py): 14 tools, 5 resources, 4 prompts.")
bullet("Assessment engine: runs the TAOSQ instrument and produces ADR, GP, ALI, ESA, and composite TAOSRI scores.")
bullet("Task classifier: assigns each task to one of five modes: automate, augment, coach, protect, hands-off.")
bullet("Interaction logger: records engagement level and skill signals (growth / stable / atrophy).")
bullet("Contrastive explanation engine: generates delta-style explanations in growth domains.")

para("Transports", bold=True)
bullet("stdio: for local MCP clients.")
bullet("Streamable HTTP + SSE: for remote MCP clients (Claude Desktop MCP Connector, future enterprise clients).")
bullet("OAuth 2.1 + PKCE (Google): protects the remote endpoint.")

para("Storage", bold=True)
bullet("Local: profiles/pro-*.md in the repository, or ~/.talent-augmenting-layer/profiles/*.md for 1-click installs. Interaction logs as JSONL.")
bullet("Hosted: PostgreSQL (Render-managed) with tables for User, Profile, AssessmentSession, ChatLog, CheckinReminder, and PilotSurvey.")

para("External Services", bold=True)
bullet("Google OAuth (authentication on both hosted and remote MCP).")
bullet("Google Drive (optional anonymised telemetry export for pilots).")
bullet("SendGrid (optional check-in reminder emails).")

doc.add_page_break()

# ==========================================================
# SECTION 3: Components in detail
# ==========================================================
h1("3. Components in Detail")

h2("3.1 The profile: the calibration layer")
para(
    "Every TAOS deployment revolves around a single artefact: a markdown profile "
    "(profiles/pro-{name}.md). The profile contains the worker's identity, their "
    "expertise map, their TAOSQ scores, a per-task-type classification table, a growth "
    "trajectory, a contrast library for contrastive explanations, and explicit red "
    "lines (tasks the AI must never do). Because it is plain markdown, it can be moved "
    "between platforms, diffed in version control, and audited by the worker."
)

h2("3.2 The TAOSQ assessment instrument")
para(
    "TAOSQ is a psychometric instrument designed specifically for AI-era workforce "
    "development. It produces five composite scores:"
)
table(
    ["Score", "Meaning"],
    [
        ["ADR: AI Dependency Risk", "How much the worker's current workflow depends on AI outputs."],
        ["GP: Growth Potential", "Where the worker has the strongest upside for skill development."],
        ["ALI: AI Literacy Index", "How well the worker understands AI strengths, limits, and failure modes."],
        ["ESA: Existing Skill Anchors", "Domains where the worker's expertise is deep enough to audit AI output."],
        ["TAOSRI: TAOS Readiness Index", "Composite readiness for AI-augmented work."],
    ],
)
para(
    "Scores are generated both by the embedded MCP assessment engine and by the hosted "
    "LLM-conversational assessment. Both paths produce identical profile files."
)

h2("3.3 The five interaction modes")
table(
    ["Mode", "AI role", "Friction", "When it applies"],
    [
        ["Automate", "Execute + annotate", "Low", "Repetitive, mechanical, well-defined work."],
        ["Augment", "Accelerate + challenge", "Low-Medium", "Complex work in the worker's expert domain."],
        ["Coach", "Scaffold + question", "Medium-High", "Skills the worker is actively developing."],
        ["Protect", "Force cognition + teach", "High", "Skills at risk of atrophying from AI over-use."],
        ["Hands-off", "Do not touch", "N/A", "Tasks that are core to the worker's human identity and judgment."],
    ],
)

h2("3.4 The MCP server and its 14 tools")
para(
    "The MCP server is the most complete integration surface. It exposes 14 tools "
    "grouped into five categories:"
)
bullet("Profile management: get_profile, get_calibration, status, list_profiles, save_profile, delete_profile.")
bullet("Assessment: assess_start, assess_score, assess_create_profile, suggest_domains.")
bullet("Runtime: classify_task, log_interaction, get_progression.")
bullet("Organisational: org_summary.")
bullet("Telemetry: parse_telemetry (extracts <tal_log> JSON blocks from LLM responses for audit and dashboarding).")

h2("3.5 Distribution channels")
para(
    "TAOS's MCP surface is shipped through four complementary distribution channels:"
)
table(
    ["Channel", "Install effort", "Profile storage"],
    [
        ["stdio MCP (mcp-server/)", "Python + config edit", "Repo or home directory"],
        ["Claude Desktop Extension (.mcpb)", "1 click", "~/.talent-augmenting-layer/"],
        ["Claude Cowork plugin (.claude-plugin)", "1 click", "~/.talent-augmenting-layer/"],
        ["Remote MCP (Streamable HTTP + OAuth)", "Google sign-in", "Hosted PostgreSQL"],
    ],
)

h2("3.6 The hosted web application")
para(
    "The hosted application at proworker-hosted.onrender.com provides a browser-based "
    "path for workers who do not use MCP-native tools. It is a FastAPI service deployed "
    "on Render with a managed PostgreSQL database. It offers:"
)
bullet("Google OAuth login with persistent per-user profiles.")
bullet("An LLM-powered conversational assessment (Gemini 2.5 Flash-Lite) that produces the same profile format as the local assessment.")
bullet("Email check-in reminders (every two weeks by default, via SendGrid).")
bullet("An anonymised daily export to Google Drive for pilot telemetry.")
bullet("The same /mcp endpoint exposed as a remote MCP server for clients that support OAuth 2.1.")

doc.add_page_break()

# ==========================================================
# SECTION 4: Data flows
# ==========================================================
h1("4. Data Flows")

h2("4.1 First-time onboarding (local)")
para("1. User runs /talent-assess in Claude Code (or equivalent in Claude Desktop).")
para("2. The MCP server calls assess_start, walks the user through the TAOSQ items conversationally, then calls assess_score.")
para("3. assess_create_profile writes a markdown profile to ~/.talent-augmenting-layer/profiles/.")
para("4. The profile is immediately available to every subsequent interaction via get_profile.")

h2("4.2 A coaching session")
para("1. User runs /talent-coach for a target skill.")
para("2. The MCP server loads the profile, classifies the task as coach, and returns the appropriate contrastive explanation scaffolding.")
para("3. Claude generates a scaffolded response using the profile's contrast library.")
para("4. log_interaction records engagement level and skill signal (growth / stable / atrophy).")
para("5. Over time, get_progression surfaces trend lines to the worker and the org dashboard.")

h2("4.3 Remote, authenticated session")
para("1. An MCP client points at proworker-hosted.onrender.com/mcp.")
para("2. The client performs the OAuth 2.1 + PKCE handshake against Google.")
para("3. The server resolves the user's profile from PostgreSQL and exposes the same 14 tools.")
para("4. All interactions are logged to the hosted ChatLog table for organisation-level analytics.")

# ==========================================================
# SECTION 5: Deployment
# ==========================================================
h1("5. Deployment & Operations")

h2("5.1 Deployed services")
table(
    ["Service", "Runtime", "Purpose"],
    [
        ["talent-augmenting-layer-hosted", "Docker on Render", "FastAPI web app + remote MCP endpoint"],
        ["PostgreSQL", "Render-managed", "Users, profiles, assessments, chat logs, reminders"],
        ["Google OAuth", "External", "Authentication for hosted + remote MCP"],
        ["SendGrid (optional)", "External", "Check-in reminder emails"],
        ["Google Drive (optional)", "External", "Anonymised pilot telemetry export"],
    ],
)

h2("5.2 Privacy and data residency")
bullet("Local tiers (1, 2, 3a-c) store profiles on the worker's own machine. No cloud, no API keys, no telemetry.")
bullet("The hosted tier stores only what is needed for assessment, progression, and reminders. All data is scoped per authenticated user.")
bullet("The Drive export is anonymised and opt-in, intended for research and pilot telemetry.")
bullet("See PRIVACY_POLICY.md for the authoritative statement.")

h2("5.3 Observability")
bullet("Structured interaction logs (JSONL locally, ChatLog table in the hosted DB).")
bullet("The <tal_log> telemetry convention lets any LLM emit machine-readable audit events that parse_telemetry can ingest.")
bullet("The org dashboard aggregates skill progression trends and atrophy warnings without exposing individual transcripts.")

# ==========================================================
# SECTION 6: Research foundation
# ==========================================================
h1("6. Research Foundation")

table(
    ["Source", "Finding", "How TAOS uses it"],
    [
        ["Buçinca et al. (2021, Harvard)", "Cognitive forcing reduces over-reliance by ~30%.", "Ask for the worker's hypothesis before revealing the AI's answer."],
        ["Buçinca et al. (2024)", "Contrastive explanations improve skills by 8% (d=0.35).", "Every growth-domain explanation uses the delta format."],
        ["Mollick et al. (2023, Wharton)", "AI: +40% quality, +26% speed: but juniors who 'just hand in' do worse.", "Protect against autopilot, especially in growth areas."],
        ["Acemoglu (MIT)", "Pro-worker AI should increase human marginal product.", "Every interaction should leave the worker more capable."],
        ["Drago & Laine (2025)", "The Intelligence Curse: humans must remain complementary to AI.", "Protect skills that maintain long-term human economic relevance."],
        ["Vygotsky, Ericsson, Deci & Ryan, Dweck", "ZPD, deliberate practice, autonomy, growth mindset.", "Calibrate scaffolding just beyond current ability; frame friction as opportunity."],
    ],
)

# ==========================================================
# SECTION 7: April 2026 Update (post-P4 deltas)
# ==========================================================
h1("7. April 2026 Update")

para(
    "This section captures the changes shipped since the 17 April 2026 revision "
    "of this document. Read it alongside the rest of the document; older sections "
    "are still substantively correct, but several behaviours, surfaces, and "
    "operational practices have moved on."
)

h2("7.1 Rebrand: Talent-Augmenting Layer is now Talent-Augmenting OS (TAOS)")
para(
    "On 30 April 2026 the product was renamed from Talent-Augmenting Layer (TAL) "
    "to Talent-Augmenting OS (TAOS). The acronym TAOS replaces TAL, the assessment "
    "instrument is TAOSQ, and the readiness index is TAOSRI. The repository slug, "
    "MCP tool prefix, and slash commands continue to use the legacy 'talent-' "
    "naming for backward compatibility with existing pilot installations; these "
    "will be migrated in a coordinated Pass 2 rebrand once the active pilot has "
    "consented to the change."
)

h2("7.2 Behaviour layer: Epistemic Rules and ambient activation")
para(
    "Three rules now override every other behavioural instruction in the system "
    "prompt, in all install channels:"
)
bullet("Calibrated confidence: do not invent facts; tag load-bearing claims with a confidence level; refuse to fabricate URLs, citations, or function signatures.")
bullet("Disagreement is a feature: do not reflexively agree; push back once with concrete counter-evidence before conceding.")
bullet("Plain voice: avoid em-dashes, AI tics, and decorative rule-of-three structures; prefer short sentences.")
para(
    "These rules are driven directly by the Vanguard pilot baseline survey, which "
    "identified hallucination, sycophancy, and generic AI voice as the three top "
    "pain points. They are present in CLAUDE.md, universal-prompt/SYSTEM_PROMPT.md, "
    "and plugin/tal-system-prompt.md."
)
para(
    "When TAOS is installed via the Claude Code plugin, a SessionStart hook "
    "(plugin/hooks/inject-tal-layer.py) prepends the TAOS system prompt and the "
    "active profile into every new session, so coaching is on from turn one without "
    "the user having to invoke a slash command."
)

h2("7.3 Speed mode and unknown-task protocol")
bullet("/talent-speed: a per-task automation override that asks the AI to skip the cognitive-forcing question and execute efficiently. Scoped to a single task; does not edit the profile; cannot override the user's red lines.")
bullet("Unknown-task protocol: when the task classifier returns 'unknown', the agent now asks the user 'automate now or get better long-term?' before proceeding, then writes the new domain and task category back to the profile with a dated change-log entry. This replaces the previous behaviour, which silently defaulted to 'augment' and let new domains drift out of the profile.")

h2("7.4 Multi-tenancy, organisation admin, and audit log")
para(
    "The data model now supports organisations and roles. Profiles, chat logs, and "
    "admin views are scoped per organisation."
)
table(
    ["Capability", "What it does"],
    [
        ["Organization and UserRole tables", "Adds org_id and role (member, admin, owner) to every user."],
        ["/admin dashboard", "Org admins see member-level skill matrix, domain coverage, risk distribution, and at-risk alerts."],
        ["Org invite flow", "Admins email-invite new members; the recipient signs in via Google and is auto-joined with the role the admin chose."],
        ["Audit log", "Append-only AuditLog table records sign-ins, role changes, invites sent and revoked, profile views, and OAuth token issuance. Available to admins as an HTML table at /admin/audit and as a CSV export at /api/admin/audit.csv."],
    ],
)

h2("7.5 Sync, install, and demo surfaces")
bullet("/talent-sync (local to hosted) and /talent-pull (hosted to local) profile sync. The /cli-token page on the hosted app issues a short-lived JWT that the CLI caches at ~/.talent-augmenting-layer/auth.json.")
bullet("Codex CLI configuration template alongside the existing Cursor and Windsurf configs.")
bullet("Public landing page and a 3-minute scripted /demo that includes both the assessment taster and a coach loop, with an explicit caveat that the coach demo is not personalised.")
bullet("5-mode compass SVG embedded in the README, the dashboard, and the landing page.")
bullet("Domain / Skill / Task glossary in CLAUDE.md and README so the three terms stop being confused.")

h2("7.6 Observability and operations")
bullet("Per-turn assessment latency: latency_ms persisted on every assessment turn, and an elapsed-seconds indicator in the assessment UI. Added so future Stan-style 20-minute stalls are visible in data, not just inferred from user complaints.")
bullet("Security headers: Content-Security-Policy, HSTS, X-Frame-Options, X-Content-Type-Options, Referrer-Policy, and Permissions-Policy on every response.")
bullet("Rate limiting via slowapi: per-IP for unauthenticated routes, per-user for authenticated routes; default 120/min, 10/min on /login.")
bullet("Render free-tier note: the hosted app sleeps after ~15 minutes of inactivity. Operationally we will move to Render Pro before any paid pilot.")

h2("7.7 Licensing and privacy")
para(
    "The licence has been changed from CC BY-NC-SA 4.0 (which made commercial use "
    "ambiguous and blocked enterprise adoption) to the Business Source License "
    "1.1, with a 4-year change date converting to Apache 2.0 on 30 April 2030. "
    "The Additional Use Grant explicitly permits internal commercial use by an "
    "organisation while reserving the right to operate a hosted or managed "
    "service to the Licensor. See LICENSE and COMMERCIAL.md for the authoritative "
    "and plain-language versions respectively."
)
para(
    "PRIVACY_POLICY.md has been rewritten as a three-channel policy (local "
    "install, hosted web app, remote MCP endpoint) with a complete subprocessor "
    "list (Render, Google OAuth, Google Gemini API, SendGrid, optional Google "
    "Drive, optional Stripe), GDPR rights, and an EU residency option for "
    "customers whose data must stay in the EU."
)

h2("7.8 Billing scaffold (gated)")
para(
    "A Stripe billing scaffold is in place behind ENABLE_BILLING=false: "
    "plan_tier and subscription_status columns on User, hosted/billing.py, "
    "and a pricing page. No code path is reachable in production while the flag "
    "is off. The scaffold exists to make a paid tier a one-flag flip when the "
    "first paying customer is signed."
)

h2("7.9 MCP correctness fixes")
bullet("Resolved three UnboundLocalError: json bugs in talent_suggest_domains, talent_assess_*, and talent_parse_telemetry caused by redundant inner 'import json' statements that shadowed the module-level import.")
bullet("Profile parser now correctly populates role, organization, industry, dependency_risk_score, growth_potential_score, skills_to_develop. Expertise extraction is scoped to Section 2 to stop contrast-library tables and the change-log divider from leaking into the parsed profile.")
bullet("Task classifier gained a 4-character prefix stemmer plus a safety bias toward protect/hands_off when classification ties.")
bullet("Markdown horizontal-rule dividers no longer leak into red-lines extraction.")

h2("7.10 Pilot context")
para(
    "The current pilot is Solita (a Finnish technology consulting firm working "
    "across multiple industries, with offices in the EU). The pilot is "
    "structured as a free proof-of-concept; the engagement that follows it, if "
    "any, will require the enterprise readiness items below."
)

h2("7.11 Open enterprise readiness items (next)")
bullet("Single sign-on via SAML and OIDC, with SCIM provisioning, ahead of any production pilot at scale.")
bullet("GDPR full-data-export and full-data-deletion endpoints surfaced in the user dashboard.")
bullet("Status page (Better Stack or equivalent) and a documented incident-response playbook with a 24-72 hour notification SLA.")
bullet("SOC 2 readiness programme via Drata or Vanta, targeting Type I in roughly 3 months and Type II in roughly 6 to 9 months.")
bullet("External penetration test before the first paying customer goes live.")
bullet("Move from Render free tier to Render Pro or equivalent paid tier before pilot users hit cold starts.")
bullet("EU data residency option offered explicitly when the customer has an EU residency requirement.")

doc.add_page_break()

# ==========================================================
# SECTION 8: Executive takeaways
# ==========================================================
h1("8. Executive Takeaways")

bullet("TAOS is a layer, not a product: it works with whichever LLM the organisation already uses.")
bullet("Profiles are portable markdown, so the worker's calibration survives platform migrations and vendor changes.")
bullet("Five interaction modes (automate, augment, coach, protect, hands-off) make the organisation's AI policy observable and measurable.")
bullet("De-skilling is a first-class risk, surfaced automatically and visible at the organisation level.")
bullet("Multiple deployment paths (1-click, plugin, remote MCP, hosted) let pilots start at the lowest possible switching cost.")
bullet("All behaviour is grounded in peer-reviewed research: defensible to regulators, boards, and workers.")

# ---------- Footer ----------
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer.add_run("Every interaction should leave the worker more capable, not more dependent.")
fr.italic = True

doc.save(OUT)
print(f"Wrote {OUT}")
